"""Generate an OpenCode Go model catalogue DOCX report."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = r"C:\Users\bohen\Documents\Hermes\Assets\opencode-go-models-use-cases.docx"

# (model, family, use case)
MODELS = [
    # MiniMax
    ("minimax-m3", "MiniMax", "Frontier flagship (428B MoE). Native multimodal (text/image/video) + 1M context. Best MiniMax pick for complex coding, agentic tasks, and multimodal reasoning."),
    ("minimax-m2.7", "MiniMax", "Recent general model. Strong all-rounder for long-document work and everyday tasks at lower cost than M3."),
    ("minimax-m2.5", "MiniMax", "Older, stable release. Reliable budget option for chat, drafting, and long-context retrieval."),
    # Kimi / Moonshot
    ("kimi-k3", "Moonshot AI (Kimi)", "Newest Kimi flagship. Top-tier multi-step reasoning, research, and agentic work."),
    ("kimi-k2.7-code", "Moonshot AI (Kimi)", "Kimi tuned for code. Code generation, debugging, refactoring, and code review."),
    ("kimi-k2.6", "Moonshot AI (Kimi)", "Strong general reasoning model. Balanced reasoning + speed for everyday complex questions."),
    ("kimi-k2.5", "Moonshot AI (Kimi)", "Proven 'thinking' model. Deep analysis, planning, and long agent loops."),
    # GLM / Zhipu
    ("glm-5.3", "Zhipu AI (GLM)", "Latest GLM flagship. Frontier general reasoning and coding."),
    ("glm-5.2", "Zhipu AI (GLM)", "Previous GLM release. Strong general-purpose tasks."),
    ("glm-5.1", "Zhipu AI (GLM)", "Earlier GLM-5 iteration. General chat, drafting, and summarization."),
    ("glm-5", "Zhipu AI (GLM)", "Base GLM-5. Dependable general assistant on the GLM stack."),
    # DeepSeek
    ("deepseek-v4-pro", "DeepSeek", "Top reasoning model (your current default). Deep research, math, and complex coding."),
    ("deepseek-v4-flash", "DeepSeek", "Fast, cheap tier. High-volume, latency-sensitive, or simple tasks."),
    ("deepseek-v4-flash-vision-exp", "DeepSeek", "Experimental vision variant of flash. Image understanding with fast reasoning."),
    # Qwen / Alibaba
    ("qwen3.8-max", "Alibaba (Qwen)", "Largest Qwen (2.4T params, 512 experts, 1M context). Frontier general/coding/agentic with huge context."),
    ("qwen3.7-max", "Alibaba (Qwen)", "Previous Qwen 'max'. Top-tier general + coding."),
    ("qwen3.7-plus", "Alibaba (Qwen)", "Balanced Qwen. Best cost/performance for everyday coding + chat."),
    ("qwen3.6-plus", "Alibaba (Qwen)", "Slightly older plus tier. Solid general workhorse."),
    ("qwen3.5-plus", "Alibaba (Qwen)", "Older plus tier. Cost-efficient general tasks and bulk text work."),
    # MiMo / Xiaomi
    ("mimo-v2.5-pro", "Xiaomi (MiMo)", "Newest MiMo pro. Strong general reasoning + coding."),
    ("mimo-v2.5", "Xiaomi (MiMo)", "Balanced MiMo. Everyday assistant tasks."),
    ("mimo-v2-pro", "Xiaomi (MiMo)", "High-end MiMo. Agentic and coding workloads."),
    ("mimo-v2-omni", "Xiaomi (MiMo)", "MiMo multimodal ('omni'). Vision and cross-modal (text + image) tasks."),
    # Frontier & specialist
    ("longcat-2.0", "LongCat AI", "1.6T MoE, open-source. Agentic coding on large codebases with 1M context."),
    ("ox-alpha-free", "OpenCode", "OpenCode's own 'OX Alpha' model on the free tier. Zero-cost general assistant + coding."),
    ("hy3", "Tencent (Hunyuan)", "Hunyuan Hy3 (295B MoE, agent-focused). Efficient agentic tasks and general reasoning."),
    ("hy3-preview", "Tencent (Hunyuan)", "Hy3 early-access build. Trying out Hy3 capabilities before the full release."),
    ("gpt-5.6-luna", "OpenAI", "GPT-5.6 'Luna' tier (lightest/cheapest variant). Cost-sensitive, high-volume workloads; 1M context."),
    ("grok-4.5", "xAI (Grok)", "Grok 4.5. General reasoning + multimodal + long-form writing."),
    ("muse-spark-1.2-contributor", "Meta (Muse Spark)", "Muse Spark 1.2 contributor tier (region-gated). Coding and tool-use tasks where available."),
]

doc = Document()

# Title
title = doc.add_heading("OpenCode Go — Model Catalogue & Use Cases", level=0)

sub = doc.add_paragraph()
r = sub.add_run("All 30 models available to your OpenCode Go subscription · retrieved live from "
                "the OpenCode Go endpoint on 25 August 2026")
r.italic = True
r.font.size = Pt(10)
r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

intro = doc.add_paragraph()
intro.add_run("OpenCode Go is the $10/month open-model tier from OpenCode (opencode.ai/zen). "
              "Each row below states what the model is best used for. Pick a model per task, not per brand.")

# Table
table = doc.add_table(rows=1, cols=3)
table.style = "Light Grid Accent 1"
hdr = table.rows[0].cells
hdr[0].text = "Model"
hdr[1].text = "Family"
hdr[2].text = "Use case"
for c in hdr:
    for p in c.paragraphs:
        for run in p.runs:
            run.bold = True

for model, fam, use in MODELS:
    row = table.add_row().cells
    row[0].text = model
    row[1].text = fam
    row[2].text = use

# Monospace the model names
for i in range(1, len(table.rows)):
    for run in table.rows[i].cells[0].paragraphs[0].runs:
        run.font.name = "Consolas"
        run.font.size = Pt(9)

# Column widths
widths = [Inches(1.7), Inches(1.5), Inches(3.9)]
for row in table.rows:
    for idx, w in enumerate(widths):
        row.cells[idx].width = w

# Quick picks
doc.add_heading("Quick picks", level=1)
picks = [
    ("Default today", "deepseek-v4-pro — deep research, math, complex coding."),
    ("Best coding", "kimi-k2.7-code, longcat-2.0, muse-spark-1.2-contributor (region-gated)."),
    ("Best agentic", "kimi-k3, hy3, qwen3.8-max."),
    ("Best multimodal", "minimax-m3, mimo-v2-omni, deepseek-v4-flash-vision-exp."),
    ("Best budget / high-volume", "gpt-5.6-luna, deepseek-v4-flash, qwen3.5-plus."),
    ("Free", "ox-alpha-free — zero-cost general assistant + coding."),
]
for k, v in picks:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(k + ": ").bold = True
    p.add_run(v)

doc.save(OUT)
print("WROTE", OUT)
print("models:", len(MODELS))
